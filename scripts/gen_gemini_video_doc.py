#!/usr/bin/env python3
"""Generate the 'How Aurion uses Gemini to interpret encounter video' .docx.

Content is sourced from the live backend code:
  - modules/providers/vision/gemini.py   (frame vs native-video paths)
  - modules/providers/vision/shared.py    (VISION_SYSTEM_PROMPT)
  - modules/vision/service.py / reconcile.py (pipeline + conflict detection)
  - modules/config/schema.py              (visual_evidence_mode, model_versions)
Run: python3 scripts/gen_gemini_video_doc.py
"""
# python-docx resolves concrete style classes (ParagraphStyle/CharacterStyle)
# at runtime; the type stubs only expose the BaseStyle base, so .font reads
# as unknown. Silence those stub-only false positives for this script.
# pyright: reportAttributeAccessIssue=false
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(0x0C, 0x1B, 0x37)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
GREY = RGBColor(0x6B, 0x72, 0x80)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
for lvl, sz in [("Heading 1", 16), ("Heading 2", 13)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = NAVY
    st.font.bold = True


def code_block(text: str):
    """Monospace, shaded paragraph for prompts / JSON."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY
    # light shading
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F3F6")
    pPr.append(shd)
    return p


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


# ---- Title ----
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("Aurion Clinical AI")
r.font.size = Pt(13)
r.font.color.rgb = GOLD
r.font.bold = True
h = doc.add_paragraph()
r = h.add_run("How Aurion Uses Gemini to Interpret Encounter Video")
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = NAVY
meta = doc.add_paragraph()
r = meta.add_run("Technical overview · MVP pilot (Descriptive Mode) · June 2026")
r.font.size = Pt(10)
r.font.color.rgb = GREY
doc.add_paragraph()

# ---- 1. Overview ----
doc.add_heading("1. Overview", level=1)
doc.add_paragraph(
    "Aurion documents a clinical encounter from three streams: audio is the spine, "
    "video is the flesh, and on-screen data is structured input. The note is built "
    "from the audio transcript first; the captured video is then used to enrich that "
    "note with what was visually observed. Google Gemini is the model Aurion uses to "
    "interpret that video."
)
doc.add_paragraph(
    "Gemini never diagnoses or interprets clinical meaning. It operates strictly in "
    "Descriptive Mode: it describes only what is literally visible (patient position, "
    "the body part being examined, observable findings, equipment, screen content) and "
    "every description is tied back to a source so it is traceable in the final note."
)

# ---- 2. When it runs ----
doc.add_heading("2. When Gemini runs (and when it does not)", level=1)
doc.add_paragraph(
    "Video interpretation is a Stage 2 step that runs only after the clinician stops "
    "recording. It is asynchronous enrichment layered on top of the Stage 1 (audio) note."
)
bullets([
    "No real-time vision. Gemini is never called during recording; live previews are audio-only. "
    "Video is interpreted only after record-stop.",
    "Audio first. Stage 1 produces the structured note from the transcript; Stage 2 adds visual "
    "evidence to it.",
    "On-device privacy gate. Frames and clips are masked on the iPhone (faces blurred, screens "
    "redacted) before anything leaves the device.",
])

# ---- 3. What is sent ----
doc.add_heading("3. What is sent to Gemini", level=1)
doc.add_paragraph(
    "An AppConfig setting, visual_evidence_mode, decides what visual evidence Aurion sends for a "
    "session. It can be changed at runtime (and overridden per session for evaluation):"
)
bullets([
    "frames_only — still JPEG frames captured at trigger timestamps (the pilot default).",
    "clips_only — short MP4 video clips around trigger moments.",
    "hybrid — clips for motion-relevant triggers (e.g. range-of-motion, gait), frames elsewhere.",
])
doc.add_paragraph("Before any media is interpreted, the backend:")
numbered([
    "Retrieves the masked frames/clips from encrypted S3 at the timestamps the trigger classifier "
    "flagged in the transcript.",
    "Validates masking. Masking is fail-closed: the iOS app guarantees only masked media is "
    "uploaded, and the service confirms the masking status in the append-only audit log before "
    "sending anything to Gemini. Raw, unmasked bytes are never sent.",
])

# ---- 4. How Gemini is called ----
doc.add_heading("4. How Gemini is called", level=1)
doc.add_paragraph(
    "Aurion calls Gemini's generateContent endpoint through its provider abstraction (never "
    "directly). There are two input paths, sharing one system prompt and one output schema:"
)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
hdr = tbl.rows[0].cells
for c, txt in zip(hdr, ["Path", "How the media is sent", "Notes"]):
    c.paragraphs[0].add_run(txt).bold = True
rows = [
    ("Still frame", "inline_data, mime image/jpeg",
     "One JPEG per trigger timestamp; the model describes that single frame."),
    ("Native video clip", "inline_data, mime video/mp4 + video_metadata.fps",
     "Gemini is the frontier model that accepts an MP4 body directly. Aurion pins the sampling "
     "rate to the configured video_capture_fps (Gemini defaults to 1 fps) so a denser clip is "
     "actually examined. The prompt asks the model to describe motion across the clip."),
]
for a, b, c in rows:
    cells = tbl.add_row().cells
    cells[0].paragraphs[0].add_run(a)
    cells[1].paragraphs[0].add_run(b)
    cells[2].paragraphs[0].add_run(c)
doc.add_paragraph()
doc.add_paragraph(
    "The Gemini model ID is config-driven (model_versions.gemini in AppConfig), so it can be "
    "upgraded without a redeploy. The pilot currently runs gemini-3.1-pro-preview; the code "
    "default is gemini-2.5-pro. Temperature (0.1) and max output tokens come from AppConfig "
    "(model_params.vision), keeping output deterministic and bounded."
)

# ---- 5. System prompt ----
doc.add_heading("5. The Descriptive-Mode system prompt", level=1)
doc.add_paragraph(
    "Every Gemini vision call carries the same system instruction. This is the single most "
    "important guardrail — it forbids diagnosis or interpretation:"
)
code_block(
    'You are a clinical visual documentation assistant. Describe only what is\n'
    'literally visible. Do not diagnose, interpret, or infer clinical meaning.\n\n'
    'Describe: patient position, visible body parts being examined, observable\n'
    'physical findings (swelling, redness, range of motion if measurable),\n'
    'equipment in use, screen content.\n'
    'Do not describe: clinical meaning, what findings suggest, what should be\n'
    'done, anything not directly visible.\n\n'
    'Return JSON only: {"description": "...", "confidence": "high|medium|low",\n'
    '"confidence_reason": "..."}\n'
    'Confidence is LOW if: blurry, wrong angle, subject not clearly visible, no\n'
    'clinically relevant content visible.'
)
doc.add_paragraph(
    "Gemini returns structured JSON: a literal description, a confidence level, and the reason "
    "for that confidence."
)

# ---- 6. From caption to note ----
doc.add_heading("6. From Gemini caption to the clinical note", level=1)
doc.add_paragraph("Each Gemini description goes through three more steps before it can touch the note:")
numbered([
    "Discard low-confidence output. Blurry, wrong-angle, or content-free captions are dropped "
    "before they can influence anything.",
    "Reconcile against the audio. Every surviving caption is compared — literally, with no "
    "clinical inference — against the audio-derived note and classified as ENRICHES, REPEATS, or "
    "CONFLICTS. This is a real comparison, not the model's self-reported guess.",
    "Merge as a traceable citation. The result is attached to the relevant note section as a "
    "frame citation, anchored to the audio segment it relates to.",
])
doc.add_paragraph("The three classifications drive different behaviour:")
bullets([
    "ENRICHES — the video adds something the audio did not capture (a location, size, laterality, "
    "or an unmentioned finding). Injected into the note.",
    "REPEATS — the video only confirms what the audio already said. Discarded so it does not "
    "duplicate the note.",
    "CONFLICTS — the video contradicts an audio claim (e.g. audio says 'no swelling' but the frame "
    "shows swelling; audio says right side, frame shows left). Flagged amber and requires mandatory "
    "physician review — 100% of conflicts must be resolved before the note can be approved.",
])

# ---- 7. Output schema ----
doc.add_heading("7. What a video citation looks like", level=1)
doc.add_paragraph(
    "Gemini's interpretation lands in the note as a structured frame-citation object. Every field "
    "keeps the description traceable and provider-attributed:"
)
code_block(
    '{\n'
    '  "frame_id":           "frame_00214",\n'
    '  "session_id":         "uuid",\n'
    '  "timestamp_ms":       14500,\n'
    '  "audio_anchor_id":    "seg_001",      // the transcript segment it relates to\n'
    '  "provider_used":      "gemini",\n'
    '  "visual_description": "...",           // Gemini\'s literal description\n'
    '  "confidence":         "high",\n'
    '  "integration_status": "ENRICHES",      // ENRICHES | REPEATS | CONFLICTS\n'
    '  "conflict_flag":      false,\n'
    '  "conflict_detail":    null\n'
    '}'
)

# ---- 8. Safety & governance ----
doc.add_heading("8. Safety and governance", level=1)
bullets([
    "Descriptive Mode only — the system prompt forbids diagnosis, interpretation, or "
    "recommendation on every single call.",
    "Fail-closed masking — Gemini only ever sees masked media whose masking status is confirmed "
    "in the audit log; raw frames never leave the device or reach the model.",
    "Full traceability — every Gemini-derived statement carries a source ID, timestamp, confidence, "
    "and provider attribution (provider_used = \"gemini\").",
    "Provider-interchangeable — Gemini is accessed through Aurion's vision-provider interface and "
    "model registry. OpenAI and Anthropic vision providers return the identical schema, so the "
    "model can be switched at runtime without changing the pipeline.",
    "No PHI in logs — only descriptive output is processed; patient identifiers are not placed in "
    "logs, errors, or storage keys.",
    "Audited and config-controlled — the model ID, evidence mode, capture rate, and token limits "
    "are AppConfig values; every change is recorded in the audit trail.",
])

# ---- 9. Configuration ----
doc.add_heading("9. Operational configuration", level=1)
doc.add_paragraph("The behaviour described above is tuned entirely through AWS AppConfig (no redeploy):")
cfg = doc.add_table(rows=1, cols=2)
cfg.style = "Light Grid Accent 1"
for c, txt in zip(cfg.rows[0].cells, ["AppConfig key", "Role"]):
    c.paragraphs[0].add_run(txt).bold = True
for k, v in [
    ("model_versions.gemini", "Which Gemini model interprets video (pilot: gemini-3.1-pro-preview)."),
    ("providers.vision", "Routes the vision stage to Gemini (vs OpenAI / Anthropic)."),
    ("pipeline.visual_evidence_mode", "frames_only / clips_only / hybrid."),
    ("pipeline.video_capture_fps", "Clip sampling rate handed to Gemini's video_metadata."),
    ("model_params.vision.temperature / max_tokens", "Determinism and output-length bounds."),
]:
    cells = cfg.add_row().cells
    cells[0].paragraphs[0].add_run(k).font.name = "Consolas"
    cells[1].paragraphs[0].add_run(v)

doc.add_paragraph()
foot = doc.add_paragraph()
r = foot.add_run(
    "Note: Gemini changes the interpretation of the video; it does not change the source of the "
    "video. In the pilot the video comes from the iPhone camera; the Ray-Ban Meta glasses path "
    "feeds the same masked-clip → Gemini pipeline once enabled."
)
r.font.size = Pt(9)
r.font.color.rgb = GREY

import os

out = os.path.normpath(
    os.path.join(os.path.dirname(__file__), "..", "docs",
                 "Aurion-Gemini-Video-Interpretation.docx")
)
doc.save(out)
print("wrote", out)
