"""Admin "Grounded Lab" endpoints — descriptive vs grounded, same clip.

The Grounded Lab is the sanctioned place to VALIDATE grounded visual findings
before trusting them on live patient notes ([[grounded-visual-findings]]). It
takes a past session whose masked frames/clips are still in S3 and re-runs the
Stage-2 vision captioning TWICE against the SAME media — once in strict
Descriptive Mode, once in Grounded Mode — then returns the two finding sets
paired by frame so a reviewer can see, side by side, exactly what grounding
adds and confirm every grounded finding is still cited to its frame.

ASYNC. Re-captioning a large frame set twice takes minutes — well past the 60s
ALB idle timeout — so a synchronous request would be dropped mid-flight even
though the server finishes. The run is therefore a background job:

  1. ``GET  /api/v1/admin/grounded-lab/sessions`` — list recent sessions that
     still have retrievable visual media. ADMIN / EVAL_TEAM.
  2. ``POST /api/v1/admin/grounded-lab/{session_id}/run`` — create a run job
     (validates the session + media synchronously, then detaches the
     captioning) and return its ``job_id`` immediately.
  3. ``GET  /api/v1/admin/grounded-lab/runs/{job_id}`` — poll the job until it
     is ``completed`` (carrying the paired result) or ``failed``.

CRITICAL — READ ONLY. Unlike ``run_stage2_vision`` (which merges captions into
a new note version and mutates the chart), this endpoint calls only the
stateless retrieve + caption steps. It NEVER writes a note version, resolves a
conflict, or touches the patient's note. Its sole side effects are the run-job
row and one PHI-free ``GROUNDED_LAB_RUN`` audit row on completion.

The comparison isolates the grounding variable: both runs use the same
template aiming (best-effort) and the same frames, differing ONLY in the
``grounded`` flag. Per-physician prompt overrides are deliberately NOT applied
here — an override would win over the grounded base prompt and mask the effect
the lab exists to measure. This means the lab shows the grounded BASE behavior,
not any one physician's customised prompt.

The lab does NOT read ``feature_flags.grounded_visual_findings_enabled``: it
always runs both modes explicitly, so it works identically whether the live
flag is on or off — you validate before flipping, and re-validate after.
"""

from __future__ import annotations

import io
import logging
import uuid
from datetime import datetime, timedelta, timezone
from typing import Literal, Optional

from botocore.exceptions import BotoCoreError, ClientError
from docx import Document
from docx.shared import Pt, RGBColor
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import Response
from pydantic import BaseModel, Field
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.v1._helpers import get_session_or_404, write_audit
from app.api.v1.admin._shared import resolve_clinician_names
from app.core.audit_events import AuditEventType
from app.core.background import spawn_background_task
from app.core.database import async_session_factory, get_db
from app.core.models import GroundedLabRunModel, SessionModel, TranscriptModel
from app.core.s3 import FRAMES_BUCKET, get_s3_client
from app.core.types import FrameCaption, Note, SessionState, Template, Transcript, UserRole
from app.modules.auth.service import CurrentUser, require_role
from app.modules.config.appconfig_client import get_config
from app.modules.config.schema import VisualEvidenceMode
from app.modules.note_gen.fusion import (
    generate_video_note,
    merge_parallel_notes,
)
from app.modules.note_gen.service import (
    get_latest_note,
    get_note_by_stage,
    resolve_session_template,
)
from app.modules.vision.reconcile import reconcile_captions
from app.modules.vision.service import (
    caption_visual_evidence,
    merge_visual_citations,
    resolve_evidence_mode,
    retrieve_all_masked_frames,
    retrieve_clips_for_triggers,
)

logger = logging.getLogger("aurion.admin.grounded_lab")

router = APIRouter(prefix="/admin", tags=["admin"])

# States in which a session's visual media has reached S3. Mirrors
# media.py's _MEDIA_BEARING_STATES — pre-review states have no frames yet,
# PURGED / failed states have nothing to replay.
_MEDIA_BEARING_STATES: frozenset[SessionState] = frozenset(
    {
        SessionState.AWAITING_REVIEW,
        SessionState.PROCESSING_STAGE2,
        SessionState.REVIEW_COMPLETE,
        SessionState.EXPORTED,
    }
)

# Paranoia ceiling on objects listed per S3 prefix (pilot clips-per-session is
# tiny). Parity with media.py so the availability signal matches.
_MAX_OBJECTS_PER_PREFIX = 500

# How far back to look for candidate sessions. Beyond the media-retention
# window the S3 lifecycle TTL has purged the frames, so a session that old
# would list with zero media anyway — bound the query to the same window.
_DEFAULT_LOOKBACK_DAYS = 30

# A run captioning a large frame set twice can legitimately take a few minutes.
# Past this budget a still-"running" job is assumed dead (its detached task
# recycled/hung) and is failed on the next poll so the UI stops spinning.
_RUN_BUDGET = timedelta(minutes=20)


# ── Schemas ──────────────────────────────────────────────────────────────────


class GroundedLabSessionItem(BaseModel):
    """One candidate session. Carries NO patient identifier."""

    session_id: str
    physician_name: str
    started_at: str
    specialty: Optional[str] = None
    visit_type: Optional[str] = None
    state: str
    frame_count: int
    clip_count: int


class GroundedLabSessionsResponse(BaseModel):
    items: list[GroundedLabSessionItem] = Field(default_factory=list)
    total: int
    page: int
    page_size: int


class GroundedLabFinding(BaseModel):
    """One captioned finding in a single mode. ``None`` when a mode produced
    no finding for that frame (e.g. discarded low-confidence)."""

    text: str
    confidence: str
    confidence_reason: str = ""
    integration_status: str
    conflict_flag: bool = False
    conflict_detail: Optional[str] = None


class GroundedLabPair(BaseModel):
    """The same frame captioned both ways, aligned for side-by-side review."""

    frame_id: str
    timestamp_ms: int
    audio_anchor_id: str
    evidence_kind: str
    descriptive: Optional[GroundedLabFinding] = None
    grounded: Optional[GroundedLabFinding] = None


class GroundedLabRunResponse(BaseModel):
    """The completed comparison payload (stored as the job's result_json)."""

    session_id: str
    specialty: Optional[str] = None
    evidence_mode: str
    provider_used: Optional[str] = None
    frame_count: int
    descriptive_findings: int
    grounded_findings: int
    pairs: list[GroundedLabPair] = Field(default_factory=list)


class GroundedLabRunStartResponse(BaseModel):
    """Returned immediately from POST /run — the job is now detached."""

    job_id: str
    session_id: str
    status: str  # always "running" here


class GroundedLabRunStatusResponse(BaseModel):
    """Poll payload. ``result`` is present iff status == "completed"."""

    job_id: str
    session_id: str
    status: str  # running | completed | failed
    error: Optional[str] = None
    result: Optional[GroundedLabRunResponse] = None


# ── Helpers ──────────────────────────────────────────────────────────────────


def _list_prefix_count(bucket: str, prefix: str) -> int:
    """Bounded object count under an S3 prefix. Failures report 0, never raise —
    availability is best-effort signal, not a hard gate."""
    client = get_s3_client()
    try:
        response = client.list_objects_v2(
            Bucket=bucket, Prefix=prefix, MaxKeys=_MAX_OBJECTS_PER_PREFIX
        )
    except (BotoCoreError, ClientError) as exc:
        logger.warning(
            "Grounded-lab media listing failed: prefix=%s: %s", prefix[:18], exc
        )
        return 0
    return len(
        [obj for obj in response.get("Contents", []) if isinstance(obj.get("Key"), str)]
    )


def _finding_from_caption(caption: FrameCaption) -> GroundedLabFinding:
    return GroundedLabFinding(
        text=caption.visual_description,
        confidence=caption.confidence,
        confidence_reason=caption.confidence_reason,
        integration_status=caption.integration_status,
        conflict_flag=caption.conflict_flag,
        conflict_detail=caption.conflict_detail,
    )


def pair_captions(
    descriptive: list[FrameCaption],
    grounded: list[FrameCaption],
) -> list[GroundedLabPair]:
    """Align two caption runs by ``frame_id`` for side-by-side display.

    Pure function (no I/O) so it is unit-testable in isolation. The union of
    frame ids is kept — a frame that a mode discarded (e.g. low confidence)
    still shows, with that side ``None``, because "grounding surfaced a finding
    where descriptive found nothing" is itself a signal worth seeing. Ordered
    by timestamp so the review reads chronologically.
    """
    desc_by_id = {c.frame_id: c for c in descriptive}
    grnd_by_id = {c.frame_id: c for c in grounded}

    # Preserve a stable timestamp/anchor per frame from whichever run has it.
    meta: dict[str, FrameCaption] = {**grnd_by_id, **desc_by_id}
    frame_ids = sorted(
        set(desc_by_id) | set(grnd_by_id),
        key=lambda fid: (meta[fid].timestamp_ms, fid),
    )

    pairs: list[GroundedLabPair] = []
    for fid in frame_ids:
        anchor = meta[fid]
        pairs.append(
            GroundedLabPair(
                frame_id=fid,
                timestamp_ms=anchor.timestamp_ms,
                audio_anchor_id=anchor.audio_anchor_id,
                evidence_kind=anchor.evidence_kind,
                descriptive=(
                    _finding_from_caption(desc_by_id[fid])
                    if fid in desc_by_id
                    else None
                ),
                grounded=(
                    _finding_from_caption(grnd_by_id[fid])
                    if fid in grnd_by_id
                    else None
                ),
            )
        )
    return pairs


async def _compute_lab_result(
    session_id: uuid.UUID, session: SessionModel, db: AsyncSession
) -> GroundedLabRunResponse:
    """Retrieve the session's evidence and caption it in BOTH modes.

    The heavy step (dozens of Gemini calls, twice). Read-only: never merges or
    persists a note. Raises ValueError if the session has no retained media.
    """
    row = (
        await db.execute(
            select(TranscriptModel).where(TranscriptModel.session_id == session_id)
        )
    ).scalar_one_or_none()
    transcript: Optional[Transcript] = (
        Transcript.model_validate_json(row.transcript_json) if row is not None else None
    )
    trigger_segments = (
        [s for s in transcript.segments if s.is_visual_trigger] if transcript else []
    )
    anchor_segments = transcript.segments if transcript else []

    evidence_mode = resolve_evidence_mode(session)

    # ALL stored frames, not just trigger-window ones: a cadence-sampled
    # session (silent exam, zero triggers) has frames in S3 that trigger
    # retrieval can't reach. Extraction already decided the set, so the lab
    # replays everything captured.
    frames = (
        await retrieve_all_masked_frames(str(session_id))
        if evidence_mode != VisualEvidenceMode.CLIPS_ONLY
        else []
    )
    clips = (
        await retrieve_clips_for_triggers(str(session_id), trigger_segments)
        if evidence_mode != VisualEvidenceMode.FRAMES_ONLY
        else []
    )
    evidence_items = [*frames, *clips]

    if not evidence_items:
        raise ValueError("no_media")

    # Template aiming (best-effort): applied to BOTH runs so the only variable
    # is the grounded flag. Never let a template failure sink the lab.
    template_for_capture: Optional[Template] = None
    if get_config().feature_flags.template_engine_enabled:
        try:
            template_for_capture = await resolve_session_template(session, db)
        except Exception:
            logger.warning(
                "Grounded-lab template resolution failed for session=%s; "
                "captioning with the base prompt",
                session_id,
                exc_info=True,
            )
            template_for_capture = None

    note: Optional[Note] = await get_latest_note(str(session_id), db)

    async def _caption(grounded: bool) -> list[FrameCaption]:
        # No per-physician prompt override: it would win over the grounded base
        # prompt and mask the very effect the lab measures.
        return await caption_visual_evidence(
            evidence=evidence_items,
            trigger_segments=trigger_segments,
            anchor_segments=anchor_segments,
            template=template_for_capture,
            note=note,
            grounded=grounded,
        )

    descriptive_caps = await _caption(grounded=False)
    grounded_caps = await _caption(grounded=True)

    pairs = pair_captions(descriptive_caps, grounded_caps)
    provider_used = next(
        (c.provider_used for c in (*grounded_caps, *descriptive_caps)), None
    )

    return GroundedLabRunResponse(
        session_id=str(session_id),
        specialty=session.specialty,
        evidence_mode=(
            evidence_mode.value if hasattr(evidence_mode, "value") else str(evidence_mode)
        ),
        provider_used=provider_used,
        frame_count=len(evidence_items),
        descriptive_findings=len(descriptive_caps),
        grounded_findings=len(grounded_caps),
        pairs=pairs,
    )


async def _run_grounded_lab_in_background(
    session_id: uuid.UUID, job_id: uuid.UUID, actor_id: uuid.UUID
) -> None:
    """Detached task: caption both modes, store the result on the job row.

    Owns its own DB session (the request that scheduled it has returned).
    Failures are recorded on the job row and never bubble. Emits the PHI-free
    ``GROUNDED_LAB_RUN`` audit on success (counts only — no caption bodies).
    """
    async with async_session_factory() as db:
        job = await db.get(GroundedLabRunModel, job_id)
        if job is None:
            return
        try:
            session = await get_session_or_404(db, session_id)
            result = await _compute_lab_result(session_id, session, db)
            job.status = "completed"
            job.result_json = result.model_dump()
            job.completed_at = datetime.now(timezone.utc)
            await db.commit()
            await write_audit(
                session_id,
                AuditEventType.GROUNDED_LAB_RUN,
                actor_id=str(actor_id),
                frame_count=result.frame_count,
                descriptive_findings=result.descriptive_findings,
                grounded_findings=result.grounded_findings,
            )
        except Exception as exc:  # noqa: BLE001 — record + swallow (detached task)
            reason = "no_media" if isinstance(exc, ValueError) else "run_failed"
            logger.warning(
                "Grounded-lab run failed: session=%s job=%s reason=%s",
                session_id,
                job_id,
                reason,
                exc_info=not isinstance(exc, ValueError),
            )
            job.status = "failed"
            job.error_message = reason
            job.completed_at = datetime.now(timezone.utc)
            await db.commit()


# ── Endpoints ────────────────────────────────────────────────────────────────


@router.get("/grounded-lab/sessions", response_model=GroundedLabSessionsResponse)
async def list_grounded_lab_sessions(
    page: int = Query(1, ge=1),
    page_size: int = Query(50, ge=1, le=200),
    user: CurrentUser = Depends(require_role(UserRole.ADMIN, UserRole.EVAL_TEAM)),
    db: AsyncSession = Depends(get_db),
) -> GroundedLabSessionsResponse:
    """List recent sessions that still have retrievable visual media.

    ADMIN or EVAL_TEAM. Selects sessions in a media-bearing state within the
    lookback window, then reports current frame/clip availability via a bounded
    per-session S3 list. NO patient identifier is included. Sessions whose media
    has been purged report zero counts rather than being hidden, so a reviewer
    can tell "no media" apart from "no such session".
    """
    window_start = datetime.now(timezone.utc) - timedelta(days=_DEFAULT_LOOKBACK_DAYS)

    base = (
        select(SessionModel)
        .where(SessionModel.state.in_(_MEDIA_BEARING_STATES))
        .where(SessionModel.created_at >= window_start)
    )
    total = (
        await db.execute(select(func.count()).select_from(base.subquery()))
    ).scalar() or 0

    stmt = (
        base.order_by(SessionModel.created_at.desc())
        .offset((page - 1) * page_size)
        .limit(page_size)
    )
    sessions = (await db.execute(stmt)).scalars().all()

    names_by_id = await resolve_clinician_names(db, (s.clinician_id for s in sessions))

    items: list[GroundedLabSessionItem] = []
    for s in sessions:
        frame_count = _list_prefix_count(FRAMES_BUCKET, f"frames/{s.id}/")
        clip_count = _list_prefix_count(FRAMES_BUCKET, f"clips/{s.id}/")
        items.append(
            GroundedLabSessionItem(
                session_id=str(s.id),
                physician_name=names_by_id[str(s.clinician_id)],
                started_at=s.created_at.isoformat() if s.created_at else "",
                specialty=s.specialty,
                visit_type=s.consultation_type,
                state=s.state.value if hasattr(s.state, "value") else str(s.state),
                frame_count=frame_count,
                clip_count=clip_count,
            )
        )

    return GroundedLabSessionsResponse(
        items=items, total=total, page=page, page_size=page_size
    )


@router.post(
    "/grounded-lab/{session_id}/run", response_model=GroundedLabRunStartResponse
)
async def run_grounded_lab(
    session_id: uuid.UUID,
    user: CurrentUser = Depends(require_role(UserRole.ADMIN, UserRole.EVAL_TEAM)),
    db: AsyncSession = Depends(get_db),
) -> GroundedLabRunStartResponse:
    """Start an async descriptive-vs-grounded replay; returns a ``job_id``.

    ADMIN or EVAL_TEAM. Validates the session (404) and that it still has
    retained media (409) synchronously — both are fast S3/DB checks — then
    detaches the captioning (minutes for a large frame set) so the request
    returns well inside the ALB idle timeout. Poll ``/runs/{job_id}`` for the
    result. READ-ONLY: the run never mutates the chart.
    """
    await get_session_or_404(db, session_id)

    # Fast media-presence check so an empty session fails immediately with 409
    # rather than spawning a job that resolves to "no findings".
    has_media = _list_prefix_count(FRAMES_BUCKET, f"frames/{session_id}/") > 0 or (
        _list_prefix_count(FRAMES_BUCKET, f"clips/{session_id}/") > 0
    )
    if not has_media:
        raise HTTPException(
            status_code=409,
            detail=(
                "No retained visual media for this session — its frames/clips "
                "have been purged or none were captured. Nothing to compare."
            ),
        )

    job = GroundedLabRunModel(
        session_id=session_id, actor_id=user.user_id, status="running"
    )
    db.add(job)
    await db.commit()
    await db.refresh(job)

    spawn_background_task(
        _run_grounded_lab_in_background(session_id, job.id, user.user_id),
        name="grounded-lab",
    )

    return GroundedLabRunStartResponse(
        job_id=str(job.id), session_id=str(session_id), status="running"
    )


@router.get(
    "/grounded-lab/runs/{job_id}", response_model=GroundedLabRunStatusResponse
)
async def get_grounded_lab_run(
    job_id: uuid.UUID,
    user: CurrentUser = Depends(require_role(UserRole.ADMIN, UserRole.EVAL_TEAM)),
    db: AsyncSession = Depends(get_db),
) -> GroundedLabRunStatusResponse:
    """Poll a lab run. ADMIN or EVAL_TEAM.

    Returns the paired result once ``completed``. A job stuck ``running`` past
    the time budget (dead detached task) is failed here so the UI stops
    spinning — the same lazy-watchdog contract the video-import poll uses.
    """
    job = await db.get(GroundedLabRunModel, job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="No such lab run.")

    if job.status == "running" and datetime.now(timezone.utc) - job.started_at > _RUN_BUDGET:
        job.status = "failed"
        job.error_message = "timed_out"
        job.completed_at = datetime.now(timezone.utc)
        await db.commit()

    result = (
        GroundedLabRunResponse.model_validate(job.result_json)
        if job.status == "completed" and job.result_json is not None
        else None
    )
    return GroundedLabRunStatusResponse(
        job_id=str(job.id),
        session_id=str(job.session_id),
        status=job.status,
        error=job.error_message,
        result=result,
    )


# ══ Fusion A vs Fusion B comparison ══════════════════════════════════════════
#
# Runs BOTH multimodal fusion architectures on the same session and returns the
# two resulting notes side by side, so the founders can pick one empirically
# before committing (the roadmap's "implement and compare the two fusion
# approaches; report which produces the better note").
#
#   Fusion A (transcript-as-context, current): Stage-1 audio note, then the
#     video captions are reconciled and MERGED as visual claims into it.
#   Fusion B (parallel-then-merge): an INDEPENDENT note is generated from the
#     video and merged section-by-section with the audio note.
#
# READ-ONLY: builds two Note objects for review; never writes a note version.


class FusionCompareResult(BaseModel):
    session_id: str
    specialty: Optional[str] = None
    frame_count: int
    # Serialized Note payloads (Note.model_dump()) for A and B.
    note_a: dict
    note_b: dict
    sections_a: int  # populated sections in A
    sections_b: int  # populated sections in B
    conflicts_b: int  # surfaced audio/visual conflicts in B


class FusionCompareStartResponse(BaseModel):
    job_id: str
    session_id: str
    status: str  # "running"


class FusionCompareStatusResponse(BaseModel):
    job_id: str
    session_id: str
    status: str  # running | completed | failed
    error: Optional[str] = None
    result: Optional[FusionCompareResult] = None


def _populated_section_count(note: Note) -> int:
    return sum(1 for s in note.sections if s.claims)


async def _compute_fusion_compare(
    session_id: uuid.UUID, session: SessionModel, db: AsyncSession
) -> FusionCompareResult:
    """Build the Fusion A and Fusion B notes from the SAME Stage-1 base + media.

    Read-only: no note version is written. Raises ValueError("no_audio_note")
    when the session has no Stage-1 note to build from, ("no_media") when its
    frames/clips are gone.
    """
    audio_note = await get_note_by_stage(str(session_id), 1, db)
    if audio_note is None:
        audio_note = await get_latest_note(str(session_id), db)
    if audio_note is None:
        raise ValueError("no_audio_note")

    row = (
        await db.execute(
            select(TranscriptModel).where(TranscriptModel.session_id == session_id)
        )
    ).scalar_one_or_none()
    transcript: Optional[Transcript] = (
        Transcript.model_validate_json(row.transcript_json) if row is not None else None
    )
    trigger_segments = (
        [s for s in transcript.segments if s.is_visual_trigger] if transcript else []
    )
    anchor_segments = transcript.segments if transcript else []
    anchor_texts = {s.id: s.text for s in anchor_segments}

    evidence_mode = resolve_evidence_mode(session)
    frames = (
        await retrieve_all_masked_frames(str(session_id))
        if evidence_mode != VisualEvidenceMode.CLIPS_ONLY
        else []
    )
    clips = (
        await retrieve_clips_for_triggers(str(session_id), trigger_segments)
        if evidence_mode != VisualEvidenceMode.FRAMES_ONLY
        else []
    )
    evidence_items = [*frames, *clips]
    if not evidence_items:
        raise ValueError("no_media")

    template: Optional[Template] = None
    if get_config().feature_flags.template_engine_enabled:
        try:
            template = await resolve_session_template(session, db)
        except Exception:
            logger.warning(
                "Fusion-compare template resolution failed for session=%s",
                session_id, exc_info=True,
            )
            template = None

    grounded = get_config().feature_flags.grounded_visual_findings_enabled

    # ── Fusion A: caption → reconcile → merge INTO a copy of the audio note.
    captions = await caption_visual_evidence(
        evidence=evidence_items,
        trigger_segments=trigger_segments,
        anchor_segments=anchor_segments,
        template=template,
        note=audio_note,
        grounded=grounded,
    )
    captions = [c for c in captions if c.confidence != "low"]
    try:
        captions = await reconcile_captions(captions, audio_note)
    except Exception:  # noqa: BLE001 — reconcile is best-effort (as in Stage 2)
        logger.warning("Fusion-compare reconcile failed; merging unreconciled")
    note_a = merge_visual_citations(
        audio_note.model_copy(deep=True), captions, template, anchor_texts
    )

    # ── Fusion B: independent video note → parallel section-weighted merge.
    video_note = await generate_video_note(
        str(session_id), evidence_items, template, grounded=grounded
    )
    note_b = merge_parallel_notes(audio_note, video_note, template)

    return FusionCompareResult(
        session_id=str(session_id),
        specialty=session.specialty,
        frame_count=len(evidence_items),
        note_a=note_a.model_dump(),
        note_b=note_b.model_dump(),
        sections_a=_populated_section_count(note_a),
        sections_b=_populated_section_count(note_b),
        conflicts_b=sum(
            1
            for s in note_b.sections
            for c in s.claims
            if c.id.startswith("conflict_")
        ),
    )


async def _run_fusion_compare_in_background(
    session_id: uuid.UUID, job_id: uuid.UUID, actor_id: uuid.UUID
) -> None:
    """Detached task: build both fusion notes, store the result on the job row."""
    async with async_session_factory() as db:
        job = await db.get(GroundedLabRunModel, job_id)
        if job is None:
            return
        try:
            session = await get_session_or_404(db, session_id)
            result = await _compute_fusion_compare(session_id, session, db)
            job.status = "completed"
            job.result_json = result.model_dump()
            job.completed_at = datetime.now(timezone.utc)
            await db.commit()
            await write_audit(
                session_id,
                AuditEventType.FUSION_COMPARE_RUN,
                actor_id=str(actor_id),
                frame_count=result.frame_count,
                sections_a=result.sections_a,
                sections_b=result.sections_b,
                conflicts_b=result.conflicts_b,
            )
        except Exception as exc:  # noqa: BLE001 — record + swallow (detached task)
            reason = (
                str(exc.args[0])
                if isinstance(exc, ValueError) and exc.args
                else "run_failed"
            )
            logger.warning(
                "Fusion-compare failed: session=%s job=%s reason=%s",
                session_id, job_id, reason,
                exc_info=not isinstance(exc, ValueError),
            )
            job.status = "failed"
            job.error_message = reason
            job.completed_at = datetime.now(timezone.utc)
            await db.commit()


@router.post(
    "/grounded-lab/{session_id}/fusion-compare",
    response_model=FusionCompareStartResponse,
)
async def run_fusion_compare(
    session_id: uuid.UUID,
    user: CurrentUser = Depends(require_role(UserRole.ADMIN, UserRole.EVAL_TEAM)),
    db: AsyncSession = Depends(get_db),
) -> FusionCompareStartResponse:
    """Start an async Fusion A vs Fusion B comparison; returns a ``job_id``.

    ADMIN or EVAL_TEAM. Validates the session (404) + media presence (409)
    synchronously, then detaches the compute (captioning + two note builds).
    Poll ``/fusion-runs/{job_id}``. READ-ONLY: never mutates the chart.
    """
    await get_session_or_404(db, session_id)
    has_media = _list_prefix_count(FRAMES_BUCKET, f"frames/{session_id}/") > 0 or (
        _list_prefix_count(FRAMES_BUCKET, f"clips/{session_id}/") > 0
    )
    if not has_media:
        raise HTTPException(
            status_code=409,
            detail=(
                "No retained visual media for this session — nothing to fuse."
            ),
        )

    job = GroundedLabRunModel(
        session_id=session_id, actor_id=user.user_id,
        run_type="fusion_compare", status="running",
    )
    db.add(job)
    await db.commit()
    await db.refresh(job)

    spawn_background_task(
        _run_fusion_compare_in_background(session_id, job.id, user.user_id),
        name="fusion-compare",
    )
    return FusionCompareStartResponse(
        job_id=str(job.id), session_id=str(session_id), status="running"
    )


@router.get(
    "/grounded-lab/fusion-runs/{job_id}",
    response_model=FusionCompareStatusResponse,
)
async def get_fusion_compare_run(
    job_id: uuid.UUID,
    user: CurrentUser = Depends(require_role(UserRole.ADMIN, UserRole.EVAL_TEAM)),
    db: AsyncSession = Depends(get_db),
) -> FusionCompareStatusResponse:
    """Poll a fusion-compare run. ADMIN or EVAL_TEAM."""
    job = await db.get(GroundedLabRunModel, job_id)
    if job is None or job.run_type != "fusion_compare":
        raise HTTPException(status_code=404, detail="No such fusion-compare run.")

    if job.status == "running" and datetime.now(timezone.utc) - job.started_at > _RUN_BUDGET:
        job.status = "failed"
        job.error_message = "timed_out"
        job.completed_at = datetime.now(timezone.utc)
        await db.commit()

    result = (
        FusionCompareResult.model_validate(job.result_json)
        if job.status == "completed" and job.result_json is not None
        else None
    )
    return FusionCompareStatusResponse(
        job_id=str(job.id),
        session_id=str(job.session_id),
        status=job.status,
        error=job.error_message,
        result=result,
    )


# ══ Modality comparison — audio-only vs visual-only vs merged ═════════════════
#
# The roadmap's toggle test page: generate the note three ways from the same
# session and read them side by side, so it's obvious what each modality
# contributes.
#
#   Audio-only : the Stage-1 note (transcript → note). What the conversation
#     alone produces.
#   Visual-only: an independent note from the video (captions → note). What the
#     images alone produce.
#   Merged     : the standard merged note (audio + visual claims) — production
#     Fusion A behaviour.
#
# The three share ONE captioning pass (visual-only + merged both need it), so
# the media is captioned once, not twice — important under the vision rate
# limit. READ-ONLY: three Note objects for review; never writes a note version.


class ModalityCompareResult(BaseModel):
    session_id: str
    specialty: Optional[str] = None
    frame_count: int
    note_audio: dict
    note_visual: Optional[dict] = None  # None when the video yielded nothing
    note_merged: dict
    sections_audio: int
    sections_visual: int
    sections_merged: int


class ModalityCompareStartResponse(BaseModel):
    job_id: str
    session_id: str
    status: str


class ModalityCompareStatusResponse(BaseModel):
    job_id: str
    session_id: str
    status: str
    error: Optional[str] = None
    result: Optional[ModalityCompareResult] = None


async def _compute_modality_compare(
    session_id: uuid.UUID, session: SessionModel, db: AsyncSession
) -> ModalityCompareResult:
    """Build the audio-only, visual-only, and merged notes from one session.

    Read-only. Captions the media ONCE and reuses it for both the visual-only
    note and the merge. Raises ValueError("no_audio_note")/("no_media").
    """
    audio_note = await get_note_by_stage(str(session_id), 1, db)
    if audio_note is None:
        audio_note = await get_latest_note(str(session_id), db)
    if audio_note is None:
        raise ValueError("no_audio_note")

    row = (
        await db.execute(
            select(TranscriptModel).where(TranscriptModel.session_id == session_id)
        )
    ).scalar_one_or_none()
    transcript: Optional[Transcript] = (
        Transcript.model_validate_json(row.transcript_json) if row is not None else None
    )
    trigger_segments = (
        [s for s in transcript.segments if s.is_visual_trigger] if transcript else []
    )
    anchor_segments = transcript.segments if transcript else []
    anchor_texts = {s.id: s.text for s in anchor_segments}

    evidence_mode = resolve_evidence_mode(session)
    frames = (
        await retrieve_all_masked_frames(str(session_id))
        if evidence_mode != VisualEvidenceMode.CLIPS_ONLY
        else []
    )
    clips = (
        await retrieve_clips_for_triggers(str(session_id), trigger_segments)
        if evidence_mode != VisualEvidenceMode.FRAMES_ONLY
        else []
    )
    evidence_items = [*frames, *clips]
    if not evidence_items:
        raise ValueError("no_media")

    template: Optional[Template] = None
    if get_config().feature_flags.template_engine_enabled:
        try:
            template = await resolve_session_template(session, db)
        except Exception:
            logger.warning(
                "Modality-compare template resolution failed for session=%s",
                session_id, exc_info=True,
            )
            template = None

    grounded = get_config().feature_flags.grounded_visual_findings_enabled

    # ── ONE captioning pass, shared by visual-only + merged.
    captions = await caption_visual_evidence(
        evidence=evidence_items,
        trigger_segments=trigger_segments,
        anchor_segments=anchor_segments,
        template=template,
        note=audio_note,
        grounded=grounded,
    )
    captions = [c for c in captions if c.confidence != "low"]

    # Visual-only note: the video alone (reuses the captions above).
    visual_note = await generate_video_note(
        str(session_id), evidence_items, template, grounded=grounded, captions=captions,
    )

    # Merged note: audio + reconciled visual claims (production Fusion A).
    merged_captions = list(captions)
    try:
        merged_captions = await reconcile_captions(merged_captions, audio_note)
    except Exception:  # noqa: BLE001 — reconcile is best-effort
        logger.warning("Modality-compare reconcile failed; merging unreconciled")
    merged_note = merge_visual_citations(
        audio_note.model_copy(deep=True), merged_captions, template, anchor_texts
    )

    return ModalityCompareResult(
        session_id=str(session_id),
        specialty=session.specialty,
        frame_count=len(evidence_items),
        note_audio=audio_note.model_dump(),
        note_visual=visual_note.model_dump() if visual_note else None,
        note_merged=merged_note.model_dump(),
        sections_audio=_populated_section_count(audio_note),
        sections_visual=_populated_section_count(visual_note) if visual_note else 0,
        sections_merged=_populated_section_count(merged_note),
    )


async def _run_modality_compare_in_background(
    session_id: uuid.UUID, job_id: uuid.UUID, actor_id: uuid.UUID
) -> None:
    async with async_session_factory() as db:
        job = await db.get(GroundedLabRunModel, job_id)
        if job is None:
            return
        try:
            session = await get_session_or_404(db, session_id)
            result = await _compute_modality_compare(session_id, session, db)
            job.status = "completed"
            job.result_json = result.model_dump()
            job.completed_at = datetime.now(timezone.utc)
            await db.commit()
            await write_audit(
                session_id,
                AuditEventType.MODALITY_COMPARE_RUN,
                actor_id=str(actor_id),
                frame_count=result.frame_count,
                sections_audio=result.sections_audio,
                sections_visual=result.sections_visual,
                sections_merged=result.sections_merged,
            )
        except Exception as exc:  # noqa: BLE001 — record + swallow (detached task)
            reason = (
                str(exc.args[0])
                if isinstance(exc, ValueError) and exc.args
                else "run_failed"
            )
            logger.warning(
                "Modality-compare failed: session=%s job=%s reason=%s",
                session_id, job_id, reason,
                exc_info=not isinstance(exc, ValueError),
            )
            job.status = "failed"
            job.error_message = reason
            job.completed_at = datetime.now(timezone.utc)
            await db.commit()


@router.post(
    "/grounded-lab/{session_id}/modality-compare",
    response_model=ModalityCompareStartResponse,
)
async def run_modality_compare(
    session_id: uuid.UUID,
    user: CurrentUser = Depends(require_role(UserRole.ADMIN, UserRole.EVAL_TEAM)),
    db: AsyncSession = Depends(get_db),
) -> ModalityCompareStartResponse:
    """Start an async audio-only vs visual-only vs merged comparison.

    ADMIN or EVAL_TEAM. Validates the session (404) + media (409), then detaches
    the compute. Poll ``/modality-runs/{job_id}``. READ-ONLY.
    """
    await get_session_or_404(db, session_id)
    has_media = _list_prefix_count(FRAMES_BUCKET, f"frames/{session_id}/") > 0 or (
        _list_prefix_count(FRAMES_BUCKET, f"clips/{session_id}/") > 0
    )
    if not has_media:
        raise HTTPException(
            status_code=409,
            detail="No retained visual media for this session — nothing to compare.",
        )

    job = GroundedLabRunModel(
        session_id=session_id, actor_id=user.user_id,
        run_type="modality_compare", status="running",
    )
    db.add(job)
    await db.commit()
    await db.refresh(job)

    spawn_background_task(
        _run_modality_compare_in_background(session_id, job.id, user.user_id),
        name="modality-compare",
    )
    return ModalityCompareStartResponse(
        job_id=str(job.id), session_id=str(session_id), status="running"
    )


@router.get(
    "/grounded-lab/modality-runs/{job_id}",
    response_model=ModalityCompareStatusResponse,
)
async def get_modality_compare_run(
    job_id: uuid.UUID,
    user: CurrentUser = Depends(require_role(UserRole.ADMIN, UserRole.EVAL_TEAM)),
    db: AsyncSession = Depends(get_db),
) -> ModalityCompareStatusResponse:
    """Poll a modality-compare run. ADMIN or EVAL_TEAM."""
    job = await db.get(GroundedLabRunModel, job_id)
    if job is None or job.run_type != "modality_compare":
        raise HTTPException(status_code=404, detail="No such modality-compare run.")

    if job.status == "running" and datetime.now(timezone.utc) - job.started_at > _RUN_BUDGET:
        job.status = "failed"
        job.error_message = "timed_out"
        job.completed_at = datetime.now(timezone.utc)
        await db.commit()

    result = (
        ModalityCompareResult.model_validate(job.result_json)
        if job.status == "completed" and job.result_json is not None
        else None
    )
    return ModalityCompareStatusResponse(
        job_id=str(job.id),
        session_id=str(job.session_id),
        status=job.status,
        error=job.error_message,
        result=result,
    )


# ── DOCX export (#export) ────────────────────────────────────────────────────
#
# Render a COMPLETED comparison result — the exact payload the client is already
# displaying — to a Word document. Read-only: it never re-runs a comparison,
# reads media, or touches a patient note; it only formats what the run produced.

_INDIGO = RGBColor(0x0F, 0x13, 0x34)   # PeriTwin indigo (matches note export)
_GREY = RGBColor(0x99, 0x99, 0x99)
_MUTED = RGBColor(0x66, 0x66, 0x66)

_EXPORT_TITLES = {
    "grounded": "Grounded Lab — Descriptive vs Grounded",
    "fusion": "Grounded Lab — Fusion A vs B",
    "modality": "Grounded Lab — Audio / Visual / Merged",
}


class ComparisonExportRequest(BaseModel):
    """The already-computed comparison the client is showing, to render as DOCX.

    ``result`` is the raw run payload (GroundedLabRunResponse / FusionCompareResult
    / ModalityCompareResult, model_dumped). Accepted as a dict so one endpoint
    serves all three shapes without re-declaring them as request bodies.
    """

    mode: Literal["grounded", "fusion", "modality"]
    session_label: str = ""
    result: dict


def _fmt_ts(ms: int) -> str:
    total = int(ms) // 1000
    return f"{total // 60}:{total % 60:02d}"


def _note_to_docx(doc: "Document", note: Optional[dict]) -> None:
    """Render a Note-shaped dict (sections -> claims) as headings + bullets."""
    if not isinstance(note, dict):
        run = doc.add_paragraph().add_run(
            "The video produced no usable note (nothing clinically visible, or "
            "captions unavailable)."
        )
        run.italic = True
        run.font.color.rgb = _GREY
        return
    sections = [
        s
        for s in note.get("sections", [])
        if isinstance(s, dict) and s.get("claims")
    ]
    if not sections:
        run = doc.add_paragraph().add_run("No populated sections.")
        run.italic = True
        run.font.color.rgb = _GREY
        return
    for section in sections:
        heading = doc.add_heading(
            section.get("title") or section.get("id", "Section"), level=2
        )
        if heading.runs:
            heading.runs[0].font.color.rgb = _INDIGO
        for claim in section.get("claims", []):
            if not isinstance(claim, dict):
                continue
            para = doc.add_paragraph(style="List Bullet")
            text_run = para.add_run(claim.get("text", ""))
            text_run.font.size = Pt(11)
            source_type = claim.get("source_type")
            if source_type:
                cite = para.add_run(f"  [{source_type}]")
                cite.font.size = Pt(9)
                cite.italic = True
                cite.font.color.rgb = _MUTED


def _finding_paragraph(doc: "Document", label: str, finding: Optional[dict]) -> None:
    para = doc.add_paragraph()
    para.add_run(f"{label}: ").bold = True
    if isinstance(finding, dict):
        para.add_run(finding.get("text", ""))
        tag = para.add_run(
            f"  [{finding.get('confidence', '')} · "
            f"{finding.get('integration_status', '')}]"
        )
        tag.font.size = Pt(9)
        tag.italic = True
        tag.font.color.rgb = _MUTED
    else:
        none_run = para.add_run("(no finding)")
        none_run.italic = True
        none_run.font.color.rgb = _GREY


def _build_comparison_docx(mode: str, session_label: str, result: dict) -> bytes:
    """Build the DOCX bytes for one comparison result. Defensive against missing
    keys so a partial payload downgrades gracefully rather than 500ing."""
    doc = Document()

    title = doc.add_heading(_EXPORT_TITLES.get(mode, "Grounded Lab — Comparison"), level=0)
    if title.runs:
        title.runs[0].font.color.rgb = _INDIGO

    meta = doc.add_paragraph()
    meta.add_run(f"Session: {session_label or '—'}").bold = True
    meta.add_run(f"  |  Frames: {result.get('frame_count', '—')}")
    if result.get("provider_used"):
        meta.add_run(f"  |  Provider: {result['provider_used']}")
    doc.add_paragraph("")

    if mode == "grounded":
        summary = doc.add_paragraph()
        summary.add_run(
            f"Descriptive findings: {result.get('descriptive_findings', 0)}  |  "
            f"Grounded findings: {result.get('grounded_findings', 0)}  |  "
            f"Evidence mode: {result.get('evidence_mode', '—')}"
        )
        doc.add_paragraph("")
        pairs = result.get("pairs") or []
        if not pairs:
            run = doc.add_paragraph().add_run("No findings produced.")
            run.italic = True
            run.font.color.rgb = _GREY
        for pair in pairs:
            if not isinstance(pair, dict):
                continue
            heading = doc.add_heading(
                f"{_fmt_ts(pair.get('timestamp_ms', 0))}  ·  {pair.get('frame_id', '')}",
                level=2,
            )
            if heading.runs:
                heading.runs[0].font.color.rgb = _INDIGO
            _finding_paragraph(doc, "Descriptive", pair.get("descriptive"))
            _finding_paragraph(doc, "Grounded", pair.get("grounded"))
            doc.add_paragraph("")
    elif mode == "fusion":
        summary = doc.add_paragraph()
        summary.add_run(
            f"Sections A: {result.get('sections_a', 0)}  |  "
            f"Sections B: {result.get('sections_b', 0)}  |  "
            f"Conflicts B: {result.get('conflicts_b', 0)}"
        )
        doc.add_paragraph("")
        for label, key in (
            ("Fusion A — transcript-as-context", "note_a"),
            ("Fusion B — parallel-then-merge", "note_b"),
        ):
            heading = doc.add_heading(label, level=1)
            if heading.runs:
                heading.runs[0].font.color.rgb = _INDIGO
            _note_to_docx(doc, result.get(key))
            doc.add_paragraph("")
    elif mode == "modality":
        summary = doc.add_paragraph()
        summary.add_run(
            f"Sections — audio: {result.get('sections_audio', 0)}  |  "
            f"visual: {result.get('sections_visual', 0)}  |  "
            f"merged: {result.get('sections_merged', 0)}"
        )
        doc.add_paragraph("")
        for label, key in (
            ("Audio only", "note_audio"),
            ("Visual only", "note_visual"),
            ("Merged (both)", "note_merged"),
        ):
            heading = doc.add_heading(label, level=1)
            if heading.runs:
                heading.runs[0].font.color.rgb = _INDIGO
            _note_to_docx(doc, result.get(key))
            doc.add_paragraph("")

    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Generated by PeriTwin Grounded Lab — a read-only validation replay. "
        "Not a patient note."
    )
    footer_run.font.size = Pt(8)
    footer_run.italic = True
    footer_run.font.color.rgb = _GREY

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@router.post("/grounded-lab/export")
async def export_grounded_lab_comparison(
    body: ComparisonExportRequest,
    user: CurrentUser = Depends(require_role(UserRole.ADMIN, UserRole.EVAL_TEAM)),
) -> Response:
    """Render a completed Grounded-Lab comparison result to a downloadable DOCX.

    Read-only: formats the payload the caller already has on screen; never
    re-runs the comparison, reads session media, or writes a note. ADMIN /
    EVAL_TEAM only (same as the run endpoints).
    """
    docx_bytes = _build_comparison_docx(body.mode, body.session_label, body.result)
    filename = f"grounded_lab_{body.mode}.docx"
    return Response(
        content=docx_bytes,
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
