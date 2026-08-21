"""Note export -- DOCX and plain text generation.

The DOCX renders the structured Note as a polished, Aurion-themed SOAP
document (Subjective / Objective / Assessment / Plan) — navy section
bands, gold accent rules, a metadata strip, and subtle source-citation
markers that preserve Aurion's traceability differentiator.

After export, triggers the cleanup pipeline to purge raw data and
migrate eval frames. Every export event is written to the audit trail.
"""

from __future__ import annotations

import io
import logging
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.audit_events import AuditEventType
from app.core.types import Note, NoteSection
from app.modules.audit_log.service import get_audit_log_service
from app.modules.cleanup.service import (
    migrate_eval_clips,
    migrate_eval_frames,
    purge_clips,
    purge_frames,
)
from app.modules.config.appconfig_client import get_config

logger = logging.getLogger("aurion.export")

# ── Aurion brand palette (sourced from Theme.swift / tailwind.config.ts) ──
_NAVY = RGBColor(0x0F, 0x13, 0x34)        # PeriTwin indigo (export brand)
_NAVY_MID = RGBColor(0x2A, 0x44, 0x8C)    # subsection headers
_NAVY_50 = "E8EBF2"                        # light fill (metadata values)
_NAVY_HEX = "0C1B37"
_GOLD = "C9A84C"                            # accent rules
_GOLD_RGB = RGBColor(0xC9, 0xA8, 0x4C)
_AMBER = "D9941F"                           # draft / status banner
_GRAY = RGBColor(0x6B, 0x72, 0x80)        # body-secondary
_GRAY_SUB = RGBColor(0x55, 0x55, 0x55)    # subtitle
_CITE = RGBColor(0x8A, 0x91, 0x9E)        # citation markers
_INK = RGBColor(0x1A, 0x1F, 0x29)         # body text
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

_BODY_FONT = "Calibri"

# SOAP grouping — note section ids bucketed into the four SOAP headers.
_SOAP_GROUPS: list[tuple[str, str, tuple[str, ...]]] = [
    (
        "S",
        "SUBJECTIVE",
        (
            "chief_complaint",
            "hpi",
            "history",
            "past_medical_history",
            "past_surgical_history",
            "medications",
            "allergies",
        ),
    ),
    (
        "O",
        "OBJECTIVE",
        (
            "vital_signs",
            "physical_exam",
            "wound_assessment",
            "functional_assessment",
            "imaging_review",
            "investigations",
        ),
    ),
    ("A", "ASSESSMENT", ("assessment",)),
    ("P", "PLAN", ("plan", "disposition")),
]

_SECTION_TITLES = {
    "chief_complaint": "Chief Complaint",
    "hpi": "History of Present Illness",
    "past_medical_history": "Past Medical History",
    "past_surgical_history": "Past Surgical History",
    "medications": "Medications",
    "allergies": "Allergies",
    "physical_exam": "Physical Examination",
    "vital_signs": "Vital Signs",
    "wound_assessment": "Wound Assessment",
    "functional_assessment": "Functional Assessment",
    "imaging_review": "Imaging Review",
    "investigations": "Investigations",
    "assessment": "Assessment",
    "plan": "Plan",
    "disposition": "Disposition",
}

_SPECIALTY_TITLES = {
    "orthopedic_surgery": "Orthopedic Surgery",
    "plastic_surgery": "Plastic Surgery",
    "musculoskeletal": "Musculoskeletal",
    "emergency_medicine": "Emergency Medicine",
    "general": "General Medicine",
}

_SOURCE_LABELS = {
    "transcript": "transcript",
    "visual": "visual",
    "screen": "screen",
    "physician_edit": "clinician",
    "measurement": "measurement",
}

_STATUS_NOTES = {
    "not_captured": "Not captured during this encounter.",
    "pending_video": "Pending video analysis.",
    "processing_failed": "Could not be processed — clinician entry required.",
}


# ── low-level docx styling helpers (drop to OOXML for shading + borders) ──
def _shade_paragraph(paragraph, fill_hex: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _bottom_border(paragraph, color_hex: str, sz: int = 6, space: int = 2) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _shade_cell(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _clear_table_borders(table) -> None:
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0")
        borders.append(e)
    tblPr.append(borders)


def _section_title(section: NoteSection) -> str:
    if section.title:
        return section.title
    return _SECTION_TITLES.get(section.id, section.id.replace("_", " ").title())


def _group_sections(note: Note) -> list[tuple[str, str, list[NoteSection]]]:
    """Bucket the note's sections into SOAP groups (order-preserving),
    routing unknown ids by keyword and appending leftovers to Plan."""
    by_id = {s.id: s for s in note.sections}
    used: set[str] = set()
    groups: list[tuple[str, str, list[NoteSection]]] = []
    for letter, label, ids in _SOAP_GROUPS:
        members = [by_id[i] for i in ids if i in by_id]
        used.update(s.id for s in members)
        groups.append((letter, label, members))

    # Heuristic placement for any section id not in the static map.
    leftovers = [s for s in note.sections if s.id not in used]
    for s in leftovers:
        sid = s.id.lower()
        if "assess" in sid or "impression" in sid:
            idx = 2
        elif "plan" in sid or "dispo" in sid or "follow" in sid:
            idx = 3
        elif any(k in sid for k in ("exam", "imag", "vital", "investig", "objective")):
            idx = 1
        else:
            idx = 0
        groups[idx][2].append(s)

    return [(letter, label, members) for letter, label, members in groups if members]


def _add_citation(paragraph, claim) -> None:
    """Append a subtle source-citation marker after a claim's text."""
    src = _SOURCE_LABELS.get(claim.source_type, claim.source_type)
    run = paragraph.add_run(f"  [{src} · {claim.source_id}]")
    run.font.size = Pt(7.5)
    run.font.color.rgb = _CITE
    run.italic = True
    if getattr(claim, "physician_edited", False):
        edit = paragraph.add_run("  (clinician-edited)")
        edit.font.size = Pt(7.5)
        edit.font.color.rgb = _GOLD_RGB
        edit.italic = True


def _build_document(note: Note):
    doc = Document()

    # Base typography + tighter clinical margins. (`styles[...]` is typed as
    # BaseStyle in the python-docx stubs; the Normal paragraph style does
    # carry `.font` at runtime.)
    normal: Any = doc.styles["Normal"]
    normal.font.name = _BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _INK
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.9)
    sec.top_margin = sec.bottom_margin = Inches(0.75)

    specialty_title = _SPECIALTY_TITLES.get(
        note.specialty, note.specialty.replace("_", " ").title()
    )
    generated = datetime.now(timezone.utc).strftime("%d %B %Y")

    # ── Masthead ──────────────────────────────────────────────────────
    eyebrow = doc.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(0)
    er = eyebrow.add_run("P E R I T W I N")
    er.font.size = Pt(8)
    er.bold = True
    er.font.color.rgb = _GOLD_RGB

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(0)
    tr = title.add_run("SOAP NOTE")
    tr.font.size = Pt(22)
    tr.bold = True
    tr.font.color.rgb = _NAVY

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(4)
    sr = subtitle.add_run(f"{specialty_title} Consultation")
    sr.font.size = Pt(11)
    sr.font.color.rgb = _GRAY_SUB

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after = Pt(6)
    _bottom_border(rule, _GOLD, sz=18, space=1)

    # ── Metadata strip (label row + value row) ────────────────────────
    meta_cols = [
        ("SPECIALTY", specialty_title),
        ("DOCUMENT", f"Stage {note.stage}  ·  Version {note.version}"),
        ("GENERATED", generated),
    ]
    table = doc.add_table(rows=2, cols=len(meta_cols))
    table.autofit = True
    _clear_table_borders(table)
    for col, (label, value) in enumerate(meta_cols):
        lcell = table.cell(0, col)
        _shade_cell(lcell, _NAVY_HEX)
        lp = lcell.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(8)
        lr.font.color.rgb = _WHITE

        vcell = table.cell(1, col)
        _shade_cell(vcell, _NAVY_50)
        vp = vcell.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(10)
        vr.font.color.rgb = _NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── SOAP body ─────────────────────────────────────────────────────
    for letter, label, sections in _group_sections(note):
        band = doc.add_paragraph()
        band.paragraph_format.space_before = Pt(10)
        band.paragraph_format.space_after = Pt(6)
        _shade_paragraph(band, _NAVY_HEX)
        br = band.add_run(f"  {letter}  —  {label}")
        br.bold = True
        br.font.size = Pt(13)
        br.font.color.rgb = _WHITE

        for section in sections:
            head = doc.add_paragraph()
            head.paragraph_format.space_before = Pt(8)
            head.paragraph_format.space_after = Pt(2)
            _bottom_border(head, _GOLD, sz=6, space=2)
            hr = head.add_run(_section_title(section))
            hr.bold = True
            hr.font.size = Pt(10.5)
            hr.font.color.rgb = _NAVY_MID

            if section.status != "populated" or not section.claims:
                note_p = doc.add_paragraph()
                note_run = note_p.add_run(
                    _STATUS_NOTES.get(section.status, "No content recorded.")
                )
                note_run.italic = True
                note_run.font.color.rgb = _GRAY
                continue

            numbered = letter == "A"
            for idx, claim in enumerate(section.claims, start=1):
                if numbered:
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Inches(0.25)
                    para.paragraph_format.space_after = Pt(3)
                    num = para.add_run(f"{idx}.  ")
                    num.bold = True
                    num.font.color.rgb = _NAVY_MID
                else:
                    para = doc.add_paragraph(style="List Bullet")
                    para.paragraph_format.space_after = Pt(2)
                body = para.add_run(claim.text)
                body.font.size = Pt(10.5)
                body.font.color.rgb = _INK
                _add_citation(para, claim)

    # ── Draft banner + provenance + descriptive-mode disclaimer ───────
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    banner = doc.add_paragraph()
    banner.paragraph_format.space_after = Pt(4)
    _shade_paragraph(banner, _AMBER)
    ban = banner.add_run(
        "  DRAFT  ·  Descriptive record — clinician review required before clinical use"
    )
    ban.bold = True
    ban.font.size = Pt(9)
    ban.font.color.rgb = _WHITE

    prov = doc.add_paragraph()
    prov.paragraph_format.space_after = Pt(0)
    pr = prov.add_run(
        f"Generated by PeriTwin  ·  Provider: {note.provider_used}  ·  "
        f"Completeness: {note.completeness_score:.0%}  ·  "
        f"Session {note.session_id}  ·  {generated}"
    )
    pr.font.size = Pt(8)
    pr.font.color.rgb = _GRAY

    disc = doc.add_paragraph()
    dr = disc.add_run(
        "This document is a descriptive record of the clinical encounter. "
        "It contains no diagnostic or interpretive conclusions."
    )
    dr.italic = True
    dr.font.size = Pt(8)
    dr.font.color.rgb = _GRAY

    return doc


async def export_note_docx(
    session_id: str,
    note: Note,
    db: AsyncSession,
) -> bytes:
    """Generate a polished, Aurion-themed SOAP DOCX from a structured Note.

    Sections are grouped into the four SOAP headers (Subjective /
    Objective / Assessment / Plan); each note section becomes a
    subsection and its claims become bulleted (or numbered, for
    Assessment) entries with subtle source-citation markers. Sections
    with status ``not_captured`` / ``pending_video`` /
    ``processing_failed`` render a status note.

    After generating the document, triggers the cleanup pipeline
    (frame/clip purge and eval migration).

    Args:
        session_id: The session being exported.
        note: The approved Note to export.
        db: Database session (passed through for any transactional needs).

    Returns:
        The DOCX file content as bytes.
    """
    audit = get_audit_log_service()

    doc = _build_document(note)

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    logger.info(
        "DOCX exported: session=%s size=%d bytes",
        session_id,
        len(docx_bytes),
    )

    # ── Audit log ─────────────────────────────────────────────────────
    await audit.write_event(
        session_id=session_id,
        event_type=AuditEventType.NOTE_EXPORTED,
        format="docx",
        version=note.version,
        stage=note.stage,
    )

    # ── Trigger cleanup pipeline ──────────────────────────────────────
    # Dual-mode evidence (P1-3): both frame and clip prefixes must be
    # migrated + purged. Each step is non-fatal — export still succeeds
    # if a cleanup leg fails (the audit row will surface the partial
    # failure to the compliance dashboard).
    try:
        await migrate_eval_frames(session_id)
    except Exception as exc:
        logger.error(
            "Eval frame migration failed during export cleanup: session=%s error=%s",
            session_id,
            str(exc),
        )

    try:
        await migrate_eval_clips(session_id)
    except Exception as exc:
        logger.error(
            "Eval clip migration failed during export cleanup: session=%s error=%s",
            session_id,
            str(exc),
        )

    # #605 — raw video (frames/clips) purge on export is spec-strict by
    # default: flag OFF (prod) purges immediately on export (stricter than the
    # 24hr spec). When ``media_review_retention_enabled`` is ON (#338), the
    # frames/clips are KEPT for the review/replay window and the S3 lifecycle
    # TTL is the max-window backstop — so we skip the immediate purge. Eval
    # migration above always runs (the eval bucket has its own retention
    # regime, independent of the clinician replay window).
    if not get_config().feature_flags.media_review_retention_enabled:
        try:
            await purge_frames(session_id)
        except Exception as exc:
            logger.error(
                "Frame purge failed during export cleanup: session=%s error=%s",
                session_id,
                str(exc),
            )

        try:
            await purge_clips(session_id)
        except Exception as exc:
            logger.error(
                "Clip purge failed during export cleanup: session=%s error=%s",
                session_id,
                str(exc),
            )

    return docx_bytes


def export_note_plaintext(session_id: str, note: Note) -> str:
    """Generate a plain text SOAP representation of a structured Note.

    This is the fallback export format when DOCX generation is not
    available (e.g., on-device iOS export). Mirrors the DOCX's SOAP
    grouping so both formats read the same.

    Args:
        session_id: The session being exported.
        note: The Note to export.

    Returns:
        The note as a plain text string.
    """
    specialty_title = _SPECIALTY_TITLES.get(
        note.specialty, note.specialty.replace("_", " ").title()
    )
    generated = datetime.now(timezone.utc).strftime("%d %B %Y")

    lines: list[str] = []
    lines.append("PERITWIN")
    lines.append("SOAP NOTE")
    lines.append(f"{specialty_title} Consultation")
    lines.append("=" * 64)
    lines.append(
        f"Specialty: {specialty_title}   |   Stage {note.stage}  v{note.version}"
        f"   |   Generated: {generated}"
    )
    lines.append("")

    for letter, label, sections in _group_sections(note):
        lines.append("")
        lines.append(f"{letter} — {label}")
        lines.append("=" * 64)
        for section in sections:
            lines.append("")
            lines.append(_section_title(section))
            lines.append("-" * 64)

            if section.status != "populated" or not section.claims:
                lines.append(
                    f"  {_STATUS_NOTES.get(section.status, 'No content recorded.')}"
                )
                continue

            numbered = letter == "A"
            for idx, claim in enumerate(section.claims, start=1):
                bullet = f"{idx}." if numbered else "•"
                src = _SOURCE_LABELS.get(claim.source_type, claim.source_type)
                lines.append(f"  {bullet} {claim.text}")
                lines.append(f"      [{src} · {claim.source_id}]")
        lines.append("")

    lines.append("=" * 64)
    lines.append("DRAFT · Descriptive record — clinician review required before clinical use")
    lines.append(
        f"Generated by PeriTwin · Provider: {note.provider_used} · "
        f"Completeness: {note.completeness_score:.0%} · Session {note.session_id}"
    )
    lines.append(
        "This document is a descriptive record of the clinical encounter. "
        "It contains no diagnostic or interpretive conclusions."
    )

    text = "\n".join(lines)

    logger.info(
        "Plaintext exported: session=%s length=%d chars",
        session_id,
        len(text),
    )

    return text
